from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import os
import json
import bcrypt
import pandas as pd
from docx import Document
from datetime import datetime
from werkzeug.utils import secure_filename
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import mimetypes
import logging
from pathlib import Path
import psycopg2
from psycopg2.extras import RealDictCursor
import psycopg2.pool

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# CORS configuration - more flexible for deployment
CORS(app, resources={
    r"/*": {
        "origins": [
            "http://127.0.0.1:5500", 
            "http://localhost:5500",
            "http://localhost:3000",
            "https://*.onrender.com",
            "https://*.vercel.app"
        ]
    }
})

# Configure MIME types for videos
mimetypes.add_type('video/mp4', '.mp4')
mimetypes.add_type('video/webm', '.webm')
mimetypes.add_type('video/ogg', '.ogv')
mimetypes.add_type('video/avi', '.avi')
mimetypes.add_type('video/mov', '.mov')
mimetypes.add_type('video/wmv', '.wmv')

# Path definitions - more robust for different environments
BASE_DIR = Path(__file__).parent.absolute()
DATA_DIR = BASE_DIR / "data"
EXPORTS_DIR = BASE_DIR / "exports"
UPLOADS_DIR = BASE_DIR / "uploads"

# Create necessary directories
DATA_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)

# File paths
USERS_FILE = DATA_DIR / "usuarios.json"
DOCUMENTS_FILE = DATA_DIR / "dados.json"
LOGS_FILE = DATA_DIR / "logs.json"
EXPORTS_FILE = DATA_DIR / "exportacoes.json"

# Allowed extensions
ALLOWED_EXTENSIONS = {
    "png", "jpg", "jpeg", "gif", "bmp",
    "mp4", "avi", "mov", "wmv",
    "mp3", "wav", "ogg", "m4a",
    "pdf", "doc", "docx", "txt"
}

# Maximum file size (16MB)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

def is_valid_file(filename):
    """Check if file has valid extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_json_file(filepath, default=None):
    """Safely load JSON file with error handling"""
    if default is None:
        default = []
    
    try:
        if filepath.exists():
            with open(filepath, "r", encoding="utf-8") as f:
                return json.load(f)
    except (json.JSONDecodeError, IOError) as e:
        logger.error(f"Error loading {filepath}: {e}")
    
    return default

def save_json_file(filepath, data):
    """Safely save JSON file with error handling"""
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        return True
    except IOError as e:
        logger.error(f"Error saving {filepath}: {e}")
        return False

def load_users():
    return load_json_file(USERS_FILE, [])

def save_users(users_list):
    return save_json_file(USERS_FILE, users_list)

def load_documents():
    return load_json_file(DOCUMENTS_FILE, [])

def save_documents(docs_list):
    return save_json_file(DOCUMENTS_FILE, docs_list)

def load_logs():
    return load_json_file(LOGS_FILE, [])

def save_logs(logs_list):
    return save_json_file(LOGS_FILE, logs_list)

def load_exports():
    return load_json_file(EXPORTS_FILE, [])

def save_exports(exports_list):
    return save_json_file(EXPORTS_FILE, exports_list)

def log_action(user, action, details=""):
    """Log user actions"""
    logs = load_logs()
    new_log = {
        "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
        "usuario": user,
        "acao": action,
        "detalhes": details,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    logs.append(new_log)
    
    # Keep only last 1000 logs to prevent file from growing too large
    if len(logs) > 1000:
        logs = logs[-1000:]
    
    save_logs(logs)
    logger.info(f"Action logged: {user} - {action}")

def can_modify(user_type):
    """Check if user can modify data"""
    return user_type in ["administrador", "editor"]

def validate_request_data(data, required_fields):
    """Validate request data has required fields"""
    if not data:
        return False, "No data provided"
    
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        return False, f"Missing required fields: {', '.join(missing_fields)}"
    
    return True, ""

# Routes
@app.route("/", methods=["GET"])
def home():
    """API documentation endpoint"""
    return jsonify({
        "name": "Sistema de Arquivologia API",
        "version": "2.0.0",
        "status": "running",
        "endpoints": {
            "authentication": ["/login", "/cadastrar_usuario"],
            "documents": ["/documentos", "/ver_dados", "/salvar_dados", "/editar_documento", "/excluir_documento"],
            "files": ["/upload_arquivo", "/uploads/<filename>", "/listar_uploads", "/excluir_upload"],
            "exports": ["/exportar_excel", "/exportar_word", "/ver_exportacoes"],
            "admin": ["/ver_usuarios", "/excluir_usuario", "/ver_logs"]
        }
    })

@app.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "directories": {
            "data": DATA_DIR.exists(),
            "exports": EXPORTS_DIR.exists(),
            "uploads": UPLOADS_DIR.exists()
        }
    })

@app.route("/login", methods=["POST"])
def login():
    """User login endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["usuario", "senha"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    users = load_users()
    username = data.get("usuario")
    password = data.get("senha", "").encode('utf-8')
    
    for user in users:
        if user["usuario"] == username:
            stored_password = user["senha"].encode('utf-8')
            if bcrypt.checkpw(password, stored_password):
                log_action(username, "LOGIN", "Login realizado com sucesso")
                return jsonify({"status": "ok", "tipo": user["tipo"]})
            else:
                log_action(username, "LOGIN_FALHOU", "Senha incorreta")
                return jsonify({"status": "erro", "mensagem": "Credenciais inválidas"}), 401
    
    log_action(username or "desconhecido", "LOGIN_FALHOU", "Usuário não encontrado")
    return jsonify({"status": "erro", "mensagem": "Credenciais inválidas"}), 401

@app.route("/cadastrar_usuario", methods=["POST"])
def register_user():
    """Register new user endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["usuario", "senha", "tipo"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    users = load_users()
    username = data.get("usuario")
    admin_user = data.get("usuario_admin", "")
    
    # Check if user already exists
    if any(u["usuario"] == username for u in users):
        return jsonify({"status": "erro", "mensagem": "Usuário já existe"}), 409
    
    # Hash password
    password_hash = bcrypt.hashpw(data.get("senha", "").encode("utf-8"), bcrypt.gensalt())
    
    new_user = {
        "usuario": username,
        "senha": password_hash.decode("utf-8"),
        "tipo": data.get("tipo"),
        "created_at": datetime.now().isoformat()
    }
    
    users.append(new_user)
    
    if save_users(users):
        log_action(admin_user, "CADASTRAR_USUARIO", f"Usuário {username} cadastrado como {data.get('tipo')}")
        return jsonify({"status": "ok", "mensagem": "Usuário cadastrado com sucesso"})
    else:
        return jsonify({"status": "erro", "mensagem": "Erro ao salvar usuário"}), 500

@app.route("/salvar_dados", methods=["POST"])
def save_document():
    """Save new document endpoint"""
    data = request.json
    
    if not data:
        return jsonify({"status": "erro", "mensagem": "Dados não fornecidos"}), 400
    
    documents = load_documents()
    user = data.get("usuario", "")
    
    # Generate ID if not provided
    if "id" not in data:
        data["id"] = datetime.now().strftime("%Y%m%d%H%M%S%f")
    
    # Set default filename if not provided
    if "arquivo_nome" not in data:
        data["arquivo_nome"] = ""
    
    # Add creation timestamp
    data["created_at"] = datetime.now().isoformat()
    
    documents.append(data)
    
    if save_documents(documents):
        log_action(user, "ADICIONAR_DOCUMENTO", f"Documento {data.get('Arquivo', '')} adicionado")
        return jsonify({"status": "ok", "mensagem": "Documento salvo com sucesso"})
    else:
        return jsonify({"status": "erro", "mensagem": "Erro ao salvar documento"}), 500

@app.route("/ver_dados", methods=["GET"])
def view_data():
    """View all documents endpoint (legacy compatibility)"""
    return jsonify(load_documents())

@app.route("/documentos", methods=["GET"])
def list_documents():
    """List all documents endpoint"""
    return jsonify(load_documents())

@app.route("/ver_usuarios", methods=["GET"])
def list_users():
    """List all users endpoint (admin only)"""
    users = load_users()
    # Remove password hashes from response
    safe_users = [{k: v for k, v in user.items() if k != "senha"} for user in users]
    return jsonify(safe_users)

@app.route("/excluir_usuario", methods=["POST"])
def delete_user():
    """Delete user endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["usuario", "tipo_usuario", "usuario_admin"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    user_type = data.get("tipo_usuario", "")
    admin_user = data.get("usuario_admin", "")
    
    if not can_modify(user_type):
        return jsonify({"status": "erro", "mensagem": "Permissão negada"}), 403
    
    users = load_users()
    username_to_delete = data.get("usuario")
    
    # Prevent self-deletion
    if username_to_delete == admin_user:
        return jsonify({"status": "erro", "mensagem": "Não é possível excluir seu próprio usuário"}), 400
    
    users = [u for u in users if u["usuario"] != username_to_delete]
    
    if save_users(users):
        log_action(admin_user, "EXCLUIR_USUARIO", f"Usuário {username_to_delete} excluído")
        return jsonify({"status": "ok", "mensagem": "Usuário excluído com sucesso"})
    else:
        return jsonify({"status": "erro", "mensagem": "Erro ao excluir usuário"}), 500

@app.route("/excluir_documento", methods=["POST"])
def delete_document():
    """Delete document endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["id", "tipo_usuario", "usuario"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    user_type = data.get("tipo_usuario", "")
    user = data.get("usuario", "")
    
    if not can_modify(user_type):
        return jsonify({"status": "erro", "mensagem": "Permissão negada"}), 403
    
    documents = load_documents()
    doc_id = data.get("id")
    
    deleted_doc = next((d for d in documents if d.get("id") == doc_id), None)
    documents = [d for d in documents if d.get("id") != doc_id]
    
    if save_documents(documents):
        if deleted_doc:
            log_action(user, "EXCLUIR_DOCUMENTO", f"Documento {deleted_doc.get('Arquivo', '')} excluído")
        return jsonify({"status": "ok", "mensagem": "Documento excluído com sucesso"})
    else:
        return jsonify({"status": "erro", "mensagem": "Erro ao excluir documento"}), 500

@app.route("/editar_documento", methods=["POST"])
def edit_document():
    """Edit document endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["original", "atualizado", "tipo_usuario", "usuario"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    user_type = data.get("tipo_usuario", "")
    user = data.get("usuario", "")
    
    if not can_modify(user_type):
        return jsonify({"status": "erro", "mensagem": "Permissão negada"}), 403
    
    documents = load_documents()
    original = data["original"]
    updated = data["atualizado"]
    
    # Add update timestamp
    updated["updated_at"] = datetime.now().isoformat()
    
    for i, doc in enumerate(documents):
        if doc.get("id") == original.get("id"):
            documents[i] = updated
            if save_documents(documents):
                log_action(user, "EDITAR_DOCUMENTO", f"Documento {updated.get('Arquivo', '')} editado")
                return jsonify({"status": "ok", "mensagem": "Documento atualizado com sucesso"})
            else:
                return jsonify({"status": "erro", "mensagem": "Erro ao salvar documento"}), 500
    
    return jsonify({"status": "erro", "mensagem": "Documento não encontrado"}), 404

@app.route("/upload_arquivo", methods=["POST"])
def upload_file():
    """File upload endpoint"""
    if "arquivo" not in request.files:
        return jsonify({"status": "erro", "mensagem": "Nenhum arquivo enviado"}), 400
    
    file = request.files["arquivo"]
    user = request.form.get("usuario", "")
    
    if file.filename == "":
        return jsonify({"status": "erro", "mensagem": "Nome do arquivo vazio"}), 400
    
    if not is_valid_file(file.filename):
        return jsonify({"status": "erro", "mensagem": "Extensão de arquivo não permitida"}), 400
    
    filename = secure_filename(file.filename)
    file_path = UPLOADS_DIR / filename
    
    # Handle duplicate filenames
    if file_path.exists():
        name_part, ext_part = os.path.splitext(filename)
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"{name_part}_{timestamp}{ext_part}"
        file_path = UPLOADS_DIR / filename
    
    try:
        file.save(str(file_path))
        log_action(user, "UPLOAD_ARQUIVO", f"Arquivo {filename} enviado")
        return jsonify({
            "status": "ok", 
            "mensagem": "Arquivo salvo com sucesso", 
            "nome_arquivo": filename
        })
    except Exception as e:
        logger.error(f"Error saving file: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao salvar arquivo"}), 500

@app.route("/uploads/<filename>", methods=["GET"])
def download_file(filename):
    """File download endpoint"""
    try:
        file_path = UPLOADS_DIR / filename
        
        if not file_path.exists():
            return jsonify({"status": "erro", "mensagem": "Arquivo não encontrado"}), 404
        
        # Detect MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        
        return send_from_directory(
            str(UPLOADS_DIR),
            filename,
            mimetype=mime_type,
            as_attachment=False,
            conditional=True
        )
    except Exception as e:
        logger.error(f"Error serving file {filename}: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao acessar arquivo"}), 500

@app.route("/listar_uploads", methods=["GET"])
def list_uploads():
    """List uploaded files endpoint"""
    try:
        files = []
        for file_path in UPLOADS_DIR.iterdir():
            if file_path.is_file():
                stat = file_path.stat()
                files.append({
                    "nome": file_path.name,
                    "tamanho": stat.st_size,
                    "data_modificacao": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                })
        
        return jsonify(files)
    except Exception as e:
        logger.error(f"Error listing uploads: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao listar arquivos"}), 500

@app.route("/excluir_upload", methods=["POST"])
def delete_upload():
    """Delete uploaded file endpoint"""
    data = request.json
    is_valid, error_msg = validate_request_data(data, ["filename", "tipo_usuario", "usuario"])
    
    if not is_valid:
        return jsonify({"status": "erro", "mensagem": error_msg}), 400
    
    filename = data.get("filename")
    user = data.get("usuario", "")
    user_type = data.get("tipo_usuario", "")
    
    if not can_modify(user_type):
        return jsonify({"status": "erro", "mensagem": "Permissão negada"}), 403
    
    try:
        file_path = UPLOADS_DIR / filename
        if file_path.exists():
            file_path.unlink()
            log_action(user, "EXCLUIR_UPLOAD", f"Arquivo {filename} excluído")
            return jsonify({"status": "ok", "mensagem": "Arquivo excluído com sucesso"})
        else:
            return jsonify({"status": "erro", "mensagem": "Arquivo não encontrado"}), 404
    except Exception as e:
        logger.error(f"Error deleting file {filename}: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao excluir arquivo"}), 500

@app.route("/exportar_excel", methods=["POST"])
def export_excel():
    """Export to Excel endpoint"""
    data = request.get_json()
    if not data:
        return jsonify({"status": "erro", "mensagem": "Dados não fornecidos"}), 400
    
    name = request.args.get("nome", f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    user = request.args.get("usuario", "")
    
    if not name.lower().endswith('.xlsx'):
        name += ".xlsx"
    
    file_path = EXPORTS_DIR / name
    
    try:
        df = pd.DataFrame(data)
        df.to_excel(str(file_path), index=False)
        
        # Register export
        exports = load_exports()
        new_export = {
            "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
            "nome_arquivo": name,
            "tipo": "Excel",
            "usuario": user,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "quantidade_documentos": len(data)
        }
        exports.append(new_export)
        save_exports(exports)
        
        log_action(user, "EXPORTAR_EXCEL", f"Exportação Excel: {name} ({len(data)} documentos)")
        
        return send_file(str(file_path), as_attachment=True, download_name=name)
    except Exception as e:
        logger.error(f"Error exporting to Excel: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao exportar para Excel"}), 500

@app.route("/exportar_word", methods=["POST"])
def export_word():
    """Export to Word endpoint"""
    data = request.get_json()
    if not data:
        return jsonify({"status": "erro", "mensagem": "Dados não fornecidos"}), 400
    
    name = request.args.get("nome", f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    user = request.args.get("usuario", "")
    
    if not name.lower().endswith('.docx'):
        name += ".docx"
    
    file_path = EXPORTS_DIR / name
    
    try:
        doc = Document()
        
        # Configure landscape orientation
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        
        doc.add_heading('Documentos Arquivísticos', level=1)
        
        if data:
            # Create table with all documents
            fields = list(data[0].keys())
            fields = [f for f in fields if f not in ["arquivo_nome", "id", "usuario"]]
            
            table = doc.add_table(rows=1, cols=len(fields))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header
            hdr_cells = table.rows[0].cells
            for i, field in enumerate(fields):
                hdr_cells[i].text = str(field)
            
            # Data
            for item in data:
                row_cells = table.add_row().cells
                for i, field in enumerate(fields):
                    row_cells[i].text = str(item.get(field, ""))
        
        doc.save(str(file_path))
        
        # Register export
        exports = load_exports()
        new_export = {
            "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
            "nome_arquivo": name,
            "tipo": "Word",
            "usuario": user,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "quantidade_documentos": len(data)
        }
        exports.append(new_export)
        save_exports(exports)
        
        log_action(user, "EXPORTAR_WORD", f"Exportação Word: {name} ({len(data)} documentos)")
        
        return send_file(str(file_path), as_attachment=True, download_name=name)
    except Exception as e:
        logger.error(f"Error exporting to Word: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao exportar para Word"}), 500

@app.route("/ver_logs", methods=["GET"])
def view_logs():
    """View logs endpoint"""
    return jsonify(load_logs())

@app.route("/ver_exportacoes", methods=["GET"])
def view_exports():
    """View exports endpoint"""
    return jsonify(load_exports())

# Error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({"status": "erro", "mensagem": "Arquivo muito grande"}), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({"status": "erro", "mensagem": "Endpoint não encontrado"}), 404

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {e}")
    return jsonify({"status": "erro", "mensagem": "Erro interno do servidor"}), 500

if __name__ == "__main__":
    # Create default admin user if no users exist
    users = load_users()
    if not users:
        admin_password = bcrypt.hashpw("admin123".encode("utf-8"), bcrypt.gensalt())
        default_admin = {
            "usuario": "admin",
            "senha": admin_password.decode("utf-8"),
            "tipo": "administrador",
            "created_at": datetime.now().isoformat()
        }
        save_users([default_admin])
        logger.info("Default admin user created: admin/admin123")
    
    # Get port from environment or default to 5000
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV') == 'development'
    
    app.run(host='0.0.0.0', port=port, debug=debug_mode)
